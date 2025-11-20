import json
import threading
import time
import tempfile
import os
from io import BytesIO
from urllib.parse import urlencode

import requests
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from PIL import Image
import tkinter as tk
from tkinter import ttk, filedialog, messagebox

GEMINI_API_ENDPOINT = "https://api.gemini.example/v1/generate" 
GEMINI_API_KEY = "YOUR_GEMINI_API_KEY"

GOOGLE_API_KEY = "YOUR_GOOGLE_API_KEY"
GOOGLE_CX = "YOUR_CUSTOM_SEARCH_ENGINE_ID"

USE_MOCK = False  
MAX_SEARCH_RESULTS = 3

def with_backoff(func, *args, max_retries=5, initial_delay=1.0, factor=2.0, **kwargs):
    """
    Выполняет функцию с экспоненциальной задержкой в случае ошибок.
    func должен бросать исключение при неудаче.
    """
    delay = initial_delay
    for attempt in range(1, max_retries + 1):
        try:
            return func(*args, **kwargs)
        except Exception as e:
            last_exc = e
            if attempt == max_retries:
                raise
            time.sleep(delay)
            delay *= factor
    raise last_exc

def gemini_call_generate_plan(topic):
    """
    Вызов Gemini для генерации Structured Output (JSON plan).
    Требуем получить валидный JSON с ключами articleTitle и researchSteps.
    """
    if USE_MOCK:
        mock = {
            "articleTitle": f"Статья: {topic}",
            "researchSteps": [
                {"step": 1, "query": f"Полное имя и биография {topic}"},
                {"step": 2, "query": f"Ключевые достижения и реформы {topic} с 2019 года"},
                {"step": 3, "query": f"Международные инициативы, связанные с {topic}"},
                {"step": 4, "query": f"Официальная фотография {topic}"}
            ]
        }
        return mock

    prompt_system = (
        "You are a professional editor. Produce a Structured Output JSON only. "
        "Output schema: {\"articleTitle\": string, \"researchSteps\": [{\"step\": int, \"query\": string}, ...]}.\n"
        f"Topic: {topic}\n"
        "Return only JSON."
    )
    headers = {"Authorization": f"Bearer {GEMINI_API_KEY}", "Content-Type": "application/json"}
    payload = {"prompt": prompt_system, "model": "gemini-2.5-flash", "max_output_tokens": 800}

    def _call():
        resp = requests.post(GEMINI_API_ENDPOINT, headers=headers, json=payload, timeout=30)
        resp.raise_for_status()
        data = resp.json()
        if isinstance(data, dict) and "output" in data:
            raw = data["output"]
        else:
            raw = json.dumps(data)
        try:
            parsed = json.loads(raw)
            return parsed
        except Exception:
            import re
            m = re.search(r'(\{.*\})', raw, flags=re.S)
            if m:
                return json.loads(m.group(1))
            else:
                raise ValueError("Не удалось распарсить JSON-план от Gemini: " + str(raw))

    return with_backoff(_call)

def gemini_call_generate_article(aggregated_data):
    """
    Вызов Gemini для финальной генерации статьи.
    aggregated_data — dict/str с собранными фактами и ссылкой на изображение.
    Ожидаем в ответе структурированный текст (включая заголовки и разделение на абзацы).
    """
    if USE_MOCK:
        article = {
            "title": aggregated_data.get("articleTitle", "Тестовая статья"),
            "sections": [
                {"heading": "Введение", "body": f"Краткое введение по теме: {aggregated_data.get('topic', '')}."},
                {"heading": "Основные факты", "body": "Собранные факты:\n" + "\n".join(aggregated_data.get("facts_summary", []))},
                {"heading": "Вывод", "body": "Короткий вывод."},
                {"heading": "ImagePlaceholder", "body": aggregated_data.get("image_url", "")}
            ]
        }
        return article

    system_instr = (
        "Using the provided data, write a coherent article divided into Introduction, headings and paragraphs. "
        "Include an explicit placeholder for the image at the end using the image URL. "
        "Return JSON: {\"title\":..., \"sections\": [{\"heading\":...,\"body\":...}, ...]}."
    )
    payload = {
        "model": "gemini-2.5-flash",
        "prompt": system_instr + "\n\nDATA:\n" + json.dumps(aggregated_data, ensure_ascii=False),
        "max_output_tokens": 1500
    }
    headers = {"Authorization": f"Bearer {GEMINI_API_KEY}", "Content-Type": "application/json"}

    def _call():
        resp = requests.post(GEMINI_API_ENDPOINT, headers=headers, json=payload, timeout=60)
        resp.raise_for_status()
        data = resp.json()
        if isinstance(data, dict) and "output" in data:
            raw = data["output"]
        else:
            raw = json.dumps(data)
        try:
            parsed = json.loads(raw)
            return parsed
        except Exception:
            import re
            m = re.search(r'(\{.*\})', raw, flags=re.S)
            if m:
                return json.loads(m.group(1))
            else:
                raise ValueError("Не удалось распарсить JSON-статью от Gemini: " + str(raw))

    return with_backoff(_call)

def google_search_text(query, api_key=GOOGLE_API_KEY, cx=GOOGLE_CX, num=3):
    """
    Возвращает список результатов: каждый результат — dict {title, snippet, link}
    """
    if USE_MOCK:
        return [{"title": f"Result for {query}", "snippet": f"Snippet about {query}", "link": f"https://example.com/{query.replace(' ','_')}"}]

    def _call():
        params = {"key": api_key, "cx": cx, "q": query, "num": min(num, 10)}
        url = "https://www.googleapis.com/customsearch/v1?" + urlencode(params)
        r = requests.get(url, timeout=15)
        r.raise_for_status()
        data = r.json()
        results = []
        for item in data.get("items", [])[:num]:
            results.append({
                "title": item.get("title"),
                "snippet": item.get("snippet"),
                "link": item.get("link")
            })
        return results

    return with_backoff(_call)

def google_search_image(query, api_key=GOOGLE_API_KEY, cx=GOOGLE_CX):
    """
    Пытается получить прямой URL подходящего изображения через Custom Search (searchType=image).
    Возвращает самый релевантный imageLink или None.
    """
    if USE_MOCK:
        return "https://via.placeholder.com/800x600.png?text=Mock+Image"

    def _call():
        params = {"key": api_key, "cx": cx, "q": query, "searchType": "image", "num": 1}
        url = "https://www.googleapis.com/customsearch/v1?" + urlencode(params)
        r = requests.get(url, timeout=15)
        r.raise_for_status()
        data = r.json()
        items = data.get("items", [])
        if not items:
            return None
        first = items[0]
        if "link" in first:
            return first["link"]
        image_obj = first.get("image", {})
        return image_obj.get("thumbnailLink") or image_obj.get("contextLink")

    return with_backoff(_call)

def create_word_document(article_json, image_url, output_filename):
    """
    article_json: {"title": str, "sections": [{"heading":..., "body":...}, ...]}
    Создает docx с Times New Roman 14pt и межстрочным 1.5. Вставляет изображение.
    """
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'  
    font.size = Pt(14)
    r = style.element.rPr
    r.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    h = doc.add_heading(article_json.get("title", "Article"), level=1)
    h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for sec in article_json.get("sections", []):
        heading = sec.get("heading", "")
        body = sec.get("body", "")

        if heading and heading.lower() not in ("imageplaceholder", "image placeholder", "image"):
            doc.add_heading(heading, level=2)
        paras = [p.strip() for p in body.split("\n\n") if p.strip()]
        for p in paras:
            par = doc.add_paragraph(p)
            par.style = doc.styles['Normal']
            par_format = par.paragraph_format
            par_format.line_spacing = 1.5

    if image_url:
        try:
            resp = requests.get(image_url, timeout=20)
            resp.raise_for_status()
            img_bytes = BytesIO(resp.content)
            image = Image.open(img_bytes)
            tmpf = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
            image.save(tmpf, format="PNG")
            tmpf.close()
            doc.add_page_break()
            doc.add_heading("Изображение", level=2)
            doc.add_picture(tmpf.name, width=Inches(6))
            os.unlink(tmpf.name)
        except Exception as e:
            doc.add_paragraph(f"Не удалось загрузить изображение. Ссылка: {image_url}")
    else:
        doc.add_paragraph("Изображение не найдено.")

    doc.save(output_filename)
    return output_filename

def execute_ai_plan(topic, save_folder=None):
    plan = gemini_call_generate_plan(topic)
    article_title = plan.get("articleTitle", topic)
    research_steps = plan.get("researchSteps", [])

    collected = []
    image_url = None
    facts_summary = []
    for step in research_steps:
        q = step.get("query")
        if not q:
            continue
        if "фото" in q.lower() or "фотограф" in q.lower() or "image" in q.lower() or "фотограф" in q.lower():
            try:
                url = google_search_image(q)
                image_url = url
                collected.append({"query": q, "type": "image", "result": url})
            except Exception as e:
                collected.append({"query": q, "type": "image", "error": str(e)})
        else:
            try:
                results = google_search_text(q, num=MAX_SEARCH_RESULTS)
                collected.append({"query": q, "type": "text", "result": results})
                for r in results:
                    facts_summary.append(f"{r.get('title')} — {r.get('snippet')} ({r.get('link')})")
            except Exception as e:
                collected.append({"query": q, "type": "text", "error": str(e)})

    aggregated = {
        "topic": topic,
        "articleTitle": article_title,
        "collected": collected,
        "image_url": image_url,
        "facts_summary": facts_summary[:20]  # trim
    }

    article_json = gemini_call_generate_article(aggregated)

    safe_title = "".join(c for c in article_title if c.isalnum() or c in " _-").strip()
    filename = f"{safe_title or 'article'}.docx"
    if save_folder:
        os.makedirs(save_folder, exist_ok=True)
        filepath = os.path.join(save_folder, filename)
    else:
        filepath = os.path.join(os.getcwd(), filename)

    create_word_document(article_json, image_url, filepath)

    return {"filepath": filepath, "article": article_json, "image_url": image_url, "plan": plan, "collected": collected}

class CADApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Когнитивный Ассемблер Документов (КАД)")
        self.geometry("700x420")
        self.resizable(False, False)

        frm = ttk.Frame(self, padding=12)
        frm.pack(fill=tk.BOTH, expand=True)

        ttk.Label(frm, text="Тема статьи:", font=("Segoe UI", 10)).grid(column=0, row=0, sticky="w")
        self.topic_var = tk.StringVar()
        self.entry = ttk.Entry(frm, textvariable=self.topic_var, width=80)
        self.entry.grid(column=0, row=1, columnspan=3, pady=8, sticky="w")

        ttk.Label(frm, text="Папка для сохранения (опционально):", font=("Segoe UI", 9)).grid(column=0, row=2, sticky="w")
        self.save_var = tk.StringVar()
        self.save_entry = ttk.Entry(frm, textvariable=self.save_var, width=60)
        self.save_entry.grid(column=0, row=3, sticky="w")
        ttk.Button(frm, text="Выбрать...", command=self.browse_folder).grid(column=1, row=3, sticky="w")

        self.generate_btn = ttk.Button(frm, text="Сгенерировать", command=self.on_generate)
        self.generate_btn.grid(column=0, row=4, pady=12, sticky="w")

        ttk.Label(frm, text="Лог:", font=("Segoe UI", 9)).grid(column=0, row=5, sticky="w")
        self.log = tk.Text(frm, height=10, width=85, state=tk.DISABLED)
        self.log.grid(column=0, row=6, columnspan=3, pady=6)

    def browse_folder(self):
        d = filedialog.askdirectory()
        if d:
            self.save_var.set(d)

    def log_write(self, text):
        self.log.configure(state=tk.NORMAL)
        self.log.insert(tk.END, f"{text}\n")
        self.log.see(tk.END)
        self.log.configure(state=tk.DISABLED)

    def on_generate(self):
        topic = self.topic_var.get().strip()
        if not topic:
            messagebox.showwarning("Ошибка", "Введите тему статьи.")
            return
        save_folder = self.save_var.get().strip() or None
        self.generate_btn.config(state=tk.DISABLED)
        self.log_write(f"Запуск генерации по теме: {topic}")

        def worker():
            try:
                res = execute_ai_plan(topic, save_folder=save_folder)
                self.log_write(f"Готово. Файл сохранён: {res['filepath']}")
                messagebox.showinfo("Успех", f"Файл сохранён:\n{res['filepath']}")
            except Exception as e:
                self.log_write(f"Ошибка: {e}")
                messagebox.showerror("Ошибка", str(e))
            finally:
                self.generate_btn.config(state=tk.NORMAL)

        threading.Thread(target=worker, daemon=True).start()

if __name__ == "__main__":
    app.mainloop()
